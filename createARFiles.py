# This script produces word documents for each species in a BirdTrack export file containing the following:
# 1.  A species header consisting of the Species name, its scientific name, the BOU Category for the species,
#     The BTO shortcode for the species and the scottish conservation status of the species.
# 2.  Observation tables for the following seasons: Winter/Spring, Summer, Autumn/Winter.
#     See the script reformatBirdtrack.py for definitions of these seasons.
#     Each table is sorted

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR

import argparse
import pandas as pd
import numpy as np
import re
import time

def createDocument():
    document = Document()

    obj_styles = document.styles

    obj_charstyle = obj_styles.add_style('SpeciesStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(18)
    obj_font.name = 'Ariel'
    obj_font.bold = True

    obj_charstyle = obj_styles.add_style('ScientificStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Ariel'
    obj_font.italic = True

    obj_charstyle = obj_styles.add_style('BouStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Ariel'
    obj_font.bold = True

    obj_charstyle = obj_styles.add_style('BtoStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(12)
    obj_font.name = 'Ariel'
    obj_font.bold = True
    return document

def createSpeciesHeader(species_df):
    scientificName = species_df['Scientific name'].unique()
    row = reference_df.query("Scientific_name == @scientificName[0]")
    if  len(row) > 0:
        bou_cat = row.iloc[0]['BOU_category']
        bto_code = row.iloc[0]['BTO_Code']
        cons_stat = row.iloc[0]['Scotland']
        p = document.add_paragraph('')
        wp = p.add_run(species, style='SpeciesStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        p.add_run(' ')
        p.add_run(scientificName, style='ScientificStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        p.add_run('\t\t')
        if bou_cat:
            reformatted_bou = ''
            first_letter = True
            for letter in bou_cat:
                if first_letter:
                    reformatted_bou = letter
                    first_letter = False
                else:
                    if re.match(r'[BCDE]', letter):
                        reformatted_bou += ', '
                    if letter != 'F':
                        reformatted_bou += letter
            p.add_run(reformatted_bou, style='BouStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        else:
            p.add_run('na', style='BouStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        p.add_run(' / ', style='BtoStyle')
        if re.match(r"^[A-Z][A-Z\.]$", bto_code):
            p.add_run(bto_code, style='BtoStyle')
        else:
            p.add_run('na', style='BtoStyle')
        p.add_run(' / ', style='BtoStyle')
        if 'Green' in cons_stat:
            p.add_run('Green', style='BtoStyle').font.color.rgb = RGBColor(0, 176, 80)
        elif 'Amber' in cons_stat:
            p.add_run('Amber', style='BtoStyle').font.color.rgb = RGBColor(255, 192, 0)
        elif 'Red' in cons_stat:
            p.add_run('Red', style='BtoStyle').font.color.rgb = RGBColor(255, 0, 0)
        else:
            p.add_run('na', style='BtoStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1

def process_season(season, document, season_df):
    if len(season_df) > 0:
        p = document.add_paragraph('')
        p.add_run(season).bold = True
        table = document.add_table(rows=1, cols=4)
        # Create the table headings
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Count'
        hdr_cells[1].text = 'Place'
        hdr_cells[2].text = 'Date'
        hdr_cells[3].text = 'Comment'
        for index, row in season_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Count'])
            row_cells[1].text = row['Place']
            row_cells[2].text = row['Date'].strftime('%d %b')
            if row['Comment'] != 0:
                row_cells[3].text = row['Comment']

def createObsTable(species_df):
    season_df = species_df[species_df['Season'] == 'Winter/Spring']
    process_season('Winter/Spring', document, season_df)

    season_df = species_df[species_df['Season'] == 'Summer']
    process_season('Summer', document, season_df)

    season_df = species_df[species_df['Season'] == 'Autumn/Winter']
    process_season('Autumn/Winter', document, season_df)

start_time = time.time()

parser = argparse.ArgumentParser(description="Create Bird Report species headings from a reformated Birdtrack export",
                                 formatter_class=argparse.ArgumentDefaultsHelpFormatter)

parser.add_argument("-f", "--file_path", type=str, required=True, help='Filepath to the Excel file to be processed')
parser.add_argument("-s", "--sheet_name", type=str, required=True, help='name of the sheet to be updated')
parser.add_argument("-d", "--data_file_path", type=str, required=True, help='Filepath to the csv file containing reference data')

args = parser.parse_args()
config = vars(args)

df = pd.read_excel(args.file_path, converters= {'Date': pd.to_datetime}, sheet_name=args.sheet_name).fillna(value = 0)
print('Input file contains {} records'.format(len(df)))

# sort the input file by BOU order, Species and Date
df.sort_values(by=['BOU order', 'Species', 'Date'], inplace=True)
print('Input file sorted')

reference_df = pd.read_csv(args.data_file_path)

#get the list of species in the report
speciesList = df.Species.unique()

for species in speciesList:
    if not species.startswith('Unidentified'):
        print(..., end='')
        print('Processing {}'.format(species), end='\r')
        species_df = df[df['Species'] == species]
        document = createDocument()
        createSpeciesHeader(species_df)
        createObsTable(species_df)
        bouOrder = species_df['BOU order'].unique()
        reformatedSpecies = species.replace(' ', '_')
        reformatedSpecies = reformatedSpecies.replace('/','')
        reformatedSpecies = str(bouOrder[0]) + '-' + reformatedSpecies
        document.save('output/' + '{}.docx'.format(reformatedSpecies))

runTime = time.time() - start_time
convert = time.strftime("%H:%M:%S", time.gmtime(runTime))
print(..., end='')
print('Execution took {}'.format(convert))