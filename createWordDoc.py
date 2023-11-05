#!/Users/Tony/.pyenv/shims/python

# TODO
# Add BOU Category, BTO short code and Conservation Status to species header
# ?Reformat as paragraph rather than table - without comments
# ?Sumarise low counts at end of seasonal section

from docx import Document
from docx.shared import Inches
import argparse
import pandas as pd
import numpy as np

def process_season(season, document, season_df):
    if len(season_df) > 0:
        p = document.add_paragraph('')
        p.add_run(season).bold = True
        table = document.add_table(rows=1, cols=4)
        # Create the table headings
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Count'
        hdr_cells[1].text = 'Place'
        hdr_cells[2].text = 'Date'
        hdr_cells[3].text = 'Comment'
        for index, row in season_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Count'])
            row_cells[1].text = row['Place']
            row_cells[2].text = row['Date']
            if row['Comment'] != 0:
                row_cells[3].text = row['Comment']

parser = argparse.ArgumentParser(description="Create Bird Report from a reformated Birdtrack export",
                                 formatter_class=argparse.ArgumentDefaultsHelpFormatter)

parser.add_argument("-f", "--file_path", type=str, required=True, help='Filepath to the Excel file to be processed')


args = parser.parse_args()
config = vars(args)

document = Document()

df = pd.read_excel(args.file_path).fillna(value = 0)

#get the list of species in the report
speciesList = df.Species.unique()

for species in speciesList:
    print(species)
    species_df = df[df['Species'] == species]
    if not species.startswith('Unidentified'):
        scientificName = species_df['Scientific name'].unique()
        p = document.add_paragraph('')
        p.add_run(species).bold = True
        p.add_run(' ')
        p.add_run(scientificName).italic = True

        season_df = species_df[species_df['Season'] == 'Winter/Spring']
        process_season('Winter/Spring', document, season_df)

        season_df = species_df[species_df['Season'] == 'Summer']
        process_season('Summer', document, season_df)

        season_df = species_df[species_df['Season'] == 'Autumn/Winter']
        process_season('Autumn/Winter', document, season_df)

        document.add_page_break()

document.save('demo.docx')