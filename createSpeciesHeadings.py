#!/Users/Tony/.pyenv/shims/python

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR

import argparse
import pandas as pd
import numpy as np
import re

parser = argparse.ArgumentParser(description="Create Bird Report species headings from a reformated Birdtrack export",
                                 formatter_class=argparse.ArgumentDefaultsHelpFormatter)

parser.add_argument("-f", "--file_path", type=str, required=True, help='Filepath to the Excel file to be processed')
parser.add_argument("-d", "--data_file_path", type=str, required=True, help='Filepath to the csv file containing reference data')

args = parser.parse_args()
config = vars(args)

document = Document()

obj_styles = document.styles

obj_charstyle = obj_styles.add_style('SpeciesStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(18)
obj_font.name = 'Ariel'
obj_font.bold = True

obj_charstyle = obj_styles.add_style('ScientificStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(14)
obj_font.name = 'Ariel'
obj_font.italic = True

obj_charstyle = obj_styles.add_style('BouStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(12)
obj_font.name = 'Ariel'
obj_font.bold = True

obj_charstyle = obj_styles.add_style('BtoStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(12)
obj_font.name = 'Ariel'
obj_font.bold = True

df = pd.read_excel(args.file_path, sheet_name='Records#1').fillna(value = 0)

# sort the input file by BOU order
df.sort_values(by=['BOU order', 'Species'], inplace=True)

reference_df = pd.read_csv(args.data_file_path)

#get the list of species in the report
speciesList = df.Species.unique()

for species in speciesList:
    print(species)
    species_df = df[df['Species'] == species]
    scientificName = species_df['Scientific name'].unique()
    row = reference_df.query("Scientific_name == @scientificName[0]")
    if  len(row) > 0:
        bou_cat = row.iloc[0]['BOU_category']
        bto_code = row.iloc[0]['BTO_Code']
        cons_stat = row.iloc[0]['Scotland']
        p = document.add_paragraph('')
        wp = p.add_run(species, style='SpeciesStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        p.add_run(' ')
        p.add_run(scientificName, style='ScientificStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        p.add_run('\t\t')
        if bou_cat:
            reformatted_bou = ''
            first_letter = True
            for letter in bou_cat:
                if first_letter:
                    reformatted_bou = letter
                    first_letter = False
                else:
                    if re.match(r'[BCDE]', letter):
                        reformatted_bou += ', '
                    if letter != 'F':
                        reformatted_bou += letter
            p.add_run(reformatted_bou, style='BouStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        else:
            p.add_run('na', style='BouStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
        p.add_run(' / ', style='BtoStyle')
        if re.match(r"^[A-Z][A-Z\.]$", bto_code):
            p.add_run(bto_code, style='BtoStyle')
        else:
            p.add_run('na', style='BtoStyle')
        p.add_run(' / ', style='BtoStyle')
        if 'Green' in cons_stat:
            p.add_run('Green', style='BtoStyle').font.color.rgb = RGBColor(0, 176, 80)
        elif 'Amber' in cons_stat:
            p.add_run('Amber', style='BtoStyle').font.color.rgb = RGBColor(255, 192, 0)
        elif 'Red' in cons_stat:
            p.add_run('Red', style='BtoStyle').font.color.rgb = RGBColor(255, 0, 0)
        else:
            p.add_run('na', style='BtoStyle').font.color.theme_color = MSO_THEME_COLOR.ACCENT_1

document.save('Generated Species Headings.docx')